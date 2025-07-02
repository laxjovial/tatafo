# shared_tools/doc_summarizer.py

import logging
from pathlib import Path
from typing import Any, Optional
import PyPDF2
import docx
import json

# For LLM integration
# from langchain_openai import ChatOpenAI
# from langchain_community.llms import GoogleGenerativeAI # For Google models
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.output_parsers import StrOutputParser

# Import config_manager and user_manager for RBAC checks
from config.config_manager import config_manager
from utils.user_manager import get_user_tier_capability, get_current_user

logger = logging.getLogger(__name__)

# --- LLM Initialization (Lazy Loading) ---
_llm_instance = None

def _get_llm_for_summarization():
    """
    Initializes and returns the LLM instance for summarization.
    Uses configuration from config_manager.
    """
    global _llm_instance
    if _llm_instance is None:
        llm_provider = config_manager.get("llm.provider", "openai")
        model_name = config_manager.get("llm.model_name", "gpt-3.5-turbo")
        temperature = config_manager.get("llm.temperature", 0.5)
        api_key = None

        if llm_provider == "openai":
            api_key = config_manager.get_secret("openai_api_key")
            if not api_key:
                logger.error("OpenAI API key not found in secrets for summarization.")
                raise ValueError("OpenAI API key is required for summarization.")
            # from langchain_openai import ChatOpenAI # Uncomment in real setup
            # _llm_instance = ChatOpenAI(model_name=model_name, temperature=temperature, api_key=api_key)
            logger.warning("Using mock LLM for summarization. Replace with actual Langchain LLM import and instantiation.")
            class MockLLM:
                def invoke(self, prompt: str) -> Any:
                    return type('obj', (object,), {'content': f"Mock summary of the provided text. Original content length: {len(prompt.split('Document Content:')[1].strip())} characters."})()
            _llm_instance = MockLLM()
        elif llm_provider == "google":
            api_key = config_manager.get_secret("google_api_key")
            if not api_key:
                logger.error("Google API key not found in secrets for summarization.")
                raise ValueError("Google API key is required for summarization.")
            # from langchain_community.llms import GoogleGenerativeAI # Uncomment in real setup
            # _llm_instance = GoogleGenerativeAI(model=model_name, temperature=temperature, google_api_key=api_key)
            logger.warning("Using mock LLM for summarization. Replace with actual Langchain LLM import and instantiation.")
            class MockLLM:
                def invoke(self, prompt: str) -> Any:
                    return type('obj', (object,), {'content': f"Mock Google summary of the provided text. Original content length: {len(prompt.split('Document Content:')[1].strip())} characters."})()
            _llm_instance = MockLLM()
        else:
            raise ValueError(f"Unsupported LLM provider for summarization: {llm_provider}")
    return _llm_instance

# --- Document Content Extraction ---
def _extract_text_from_pdf(file_path: Path) -> str:
    """Extracts text from a PDF file."""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text() or ""
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}", exc_info=True)
        raise ValueError(f"Could not extract text from PDF: {e}")

def _extract_text_from_docx(file_path: Path) -> str:
    """Extracts text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}", exc_info=True)
        raise ValueError(f"Could not extract text from DOCX: {e}")

def _extract_text_from_txt(file_path: Path) -> str:
    """Extracts text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {e}", exc_info=True)
        raise ValueError(f"Could not extract text from TXT: {e}")

# --- Summarization Tool ---
def summarize_document(file_path: Path, user_token: str = "default") -> str:
    """
    Summarizes the content of a document located at the given file path using an LLM.
    Supports PDF, DOCX, and TXT files.
    Applies RBAC checks for summarization capabilities.

    Args:
        file_path (Path): The path to the document file.
        user_token (str, optional): The unique identifier for the user. Defaults to "default".
                                    Used for RBAC capability checks.

    Returns:
        str: A concise summary of the document content, or an error message.
    """
    logger.info(f"Attempting to summarize document: {file_path} for user: {user_token}")

    # RBAC Check for Summarization Enabled
    if not get_user_tier_capability(user_token, 'summarization_enabled', False):
        return "Error: Document summarization is not enabled for your current tier."
    
    # Get user's allowed max input characters for summarization
    max_input_chars_allowed = get_user_tier_capability(user_token, 'summarization_max_input_chars', config_manager.get('llm.max_summary_input_chars', 10000))

    if not file_path.exists():
        logger.error(f"Document not found at '{file_path}'.")
        return f"Error: Document not found at '{file_path}'."

    extracted_text = ""
    file_extension = file_path.suffix.lower()

    try:
        if file_extension == '.pdf':
            extracted_text = _extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            extracted_text = _extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            extracted_text = _extract_text_from_txt(file_path)
        else:
            return f"Error: Unsupported file type for summarization: {file_extension}. Supported types: .pdf, .docx, .txt."
        
        if not extracted_text.strip():
            return "Error: Could not extract any readable text from the document."

        # Truncate text based on user's allowed max_input_chars
        if len(extracted_text) > max_input_chars_allowed:
            logger.warning(f"Document text (length {len(extracted_text)}) truncated to {max_input_chars_allowed} characters for summarization.")
            extracted_text = extracted_text[:max_input_chars_allowed]
        
        llm = _get_llm_for_summarization()
        
        prompt_template = """
        You are an expert summarizer. Summarize the following document content concisely and accurately.
        Focus on the main points, key findings, and conclusions.

        Document Title: {document_title}
        Document Content:
        {document_content}

        Concise Summary:
        """
        
        # In a real Langchain setup, you'd use a prompt template and parser
        # prompt = ChatPromptTemplate.from_template(prompt_template)
        # chain = prompt | llm | StrOutputParser()
        # summary = chain.invoke({"document_title": file_path.name, "document_content": extracted_text})

        # For mock LLM, simulate the invocation
        mock_prompt_input = prompt_template.format(document_title=file_path.name, document_content=extracted_text)
        summary = llm.invoke(mock_prompt_input).content

        logger.info(f"Successfully summarized document: {file_path.name}")
        return summary

    except ValueError as ve:
        logger.error(f"Summarization failed due to data/extraction error: {ve}", exc_info=True)
        return f"Error during summarization: {ve}"
    except Exception as e:
        logger.critical(f"An unexpected error occurred during document summarization for {file_path}: {e}", exc_info=True)
        return f"An unexpected error occurred during summarization: {e}"

# CLI Test (optional)
if __name__ == "__main__":
    import shutil
    from unittest.mock import MagicMock
    import sys
    import os

    logging.basicConfig(level=logging.INFO)

    # Mock Streamlit secrets and config_manager for local testing
    class MockSecrets:
        def __init__(self):
            self.openai = {"api_key": "sk-mock-openai-key-12345"}
            self.google = {"api_key": "AIzaSy-mock-google-key"}
            self.user_tokens = {
                "free_user_token": "mock_free_token",
                "pro_user_token": "mock_pro_token",
                "premium_user_token": "mock_premium_token",
                "admin_user_token": "mock_admin_token"
            }
            self.firebase_config = "{}" # Mock empty config for Firebase if not set

        def get(self, key, default=None):
            parts = key.split('.')
            val = self
            for part in parts:
                if hasattr(val, part):
                    val = getattr(val, part)
                elif isinstance(val, dict) and part in val:
                    val = val[part]
                else:
                    return default
            return val
    
    class MockConfigManager:
        _instance = None
        _is_loaded = False
        def __init__(self):
            if MockConfigManager._instance is not None:
                raise Exception("ConfigManager is a singleton. Use get_instance().")
            MockConfigManager._instance = self
            self._config_data = {
                'llm': {
                    'provider': 'openai', # Mock provider
                    'model_name': 'gpt-3.5-turbo',
                    'temperature': 0.5,
                    'max_tokens': 4096,
                    'max_summary_input_chars': 10000 # Default config value
                },
                'rag': {'chunk_size': 500, 'chunk_overlap': 50, 'max_query_results_k': 10},
                'web_scraping': {
                    'user_agent': 'Mozilla/5.0 (Test; Python)',
                    'timeout_seconds': 5,
                    'max_search_results': 5
                },
                'tiers': {}, # This will be overridden by tiers.yaml
                'default_user_tier': 'free',
                'default_user_roles': ['user'],
                'api_configs': []
            }
            self._is_loaded = True
        
        def get(self, key, default=None):
            parts = key.split('.')
            val = self._config_data
            for part in parts:
                if isinstance(val, dict) and part in val:
                    val = val[part]
                else:
                    return default
            return val
        
        def get_secret(self, key, default=None):
            # Simulate returning mock API key
            if key == "openai_api_key": return "MOCK_OPENAI_KEY_123"
            if key == "google_api_key": return "MOCK_GOOGLE_KEY_456"
            return st.secrets.get(key, default)

        def set_secret(self, key, value):
            setattr(st.secrets, key, value)


    # Mock user_manager.get_current_user and get_user_tier_capability for testing RBAC
    class MockUserManager:
        _mock_users = {
            "mock_free_token": {"user_id": "mock_free_token", "username": "FreeUser", "email": "free@example.com", "tier": "free", "roles": ["user"]},
            "mock_pro_token": {"user_id": "mock_pro_token", "username": "ProUser", "email": "pro@example.com", "tier": "pro", "roles": ["user"]},
            "mock_premium_token": {"user_id": "mock_premium_token", "username": "PremiumUser", "email": "premium@example.com", "tier": "premium", "roles": ["user"]},
            "mock_admin_token": {"user_id": "mock_admin_token", "username": "AdminUser", "email": "admin@example.com", "tier": "admin", "roles": ["user", "admin"]},
        }
        _rbac_capabilities = {
            'capabilities': {
                'summarization_enabled': {
                    'default': False,
                    'roles': {'pro': True, 'premium': True, 'admin': True}
                },
                'summarization_max_input_chars': {
                    'default': 5000,
                    'roles': {'pro': 10000, 'premium': 20000, 'admin': 50000}
                }
            }
        }
        _tier_hierarchy = {
            "free": 0, "user": 1, "basic": 2, "pro": 3, "premium": 4, "admin": 99
        }

        def get_current_user(self) -> Dict[str, Any]:
            return getattr(self, '_current_mock_user', {})

        def get_user_tier_capability(self, user_token: Optional[str], capability_key: str, default_value: Any = None) -> Any:
            user_info = self._mock_users.get(user_token, {})
            user_id = user_info.get('user_id')
            user_tier = user_info.get('tier', 'free')
            user_roles = user_info.get('roles', [])

            if "admin" in user_roles:
                if isinstance(default_value, bool): return True
                if isinstance(default_value, (int, float)): return float('inf')
                return default_value
            
            capability_config = self._rbac_capabilities.get('capabilities', {}).get(capability_key)
            if not capability_config:
                return default_value

            for role in user_roles:
                if role in capability_config.get('roles', {}):
                    return capability_config['roles'][role]
            
            return capability_config.get('default', default_value)

    # Patch the actual imports for testing
    import streamlit as st_mock
    if not hasattr(st_mock, 'secrets'):
        st_mock.secrets = MockSecrets()
    
    sys.modules['config.config_manager'].config_manager = MockConfigManager()
    sys.modules['config.config_manager'].ConfigManager = MockConfigManager
    sys.modules['utils.user_manager'] = MockUserManager()
    sys.modules['utils.user_manager']._RBAC_CAPABILITIES = MockUserManager()._rbac_capabilities
    sys.modules['utils.user_manager']._TIER_HIERARCHY = MockUserManager()._tier_hierarchy

    # Reset _llm_instance for each test run
    global _llm_instance
    _llm_instance = None

    # Create dummy files for testing
    test_dir = Path("./test_docs_for_summarizer")
    test_dir.mkdir(exist_ok=True)

    pdf_path = test_dir / "sample.pdf"
    docx_path = test_dir / "sample.docx"
    txt_path = test_dir / "sample.txt"
    empty_txt_path = test_dir / "empty.txt"
    long_txt_path = test_dir / "long_sample.txt"

    # Create dummy PDF (requires PyPDF2 to write, so just create a simple one)
    # For a real test, you'd need a pre-existing PDF or a more complex generation.
    # Here, we'll mock PyPDF2 for simplicity in testing text extraction.
    class MockPdfReader:
        def __init__(self, text_content):
            self.pages = [MagicMock()]
            self.pages[0].extract_text.return_value = text_content
        def __len__(self): return 1
    
    # Create dummy DOCX
    doc = docx.Document()
    doc.add_paragraph("This is a sample DOCX document for summarization testing.")
    doc.add_paragraph("It contains a few sentences about various topics.")
    doc.save(docx_path)

    # Create dummy TXT
    with open(txt_path, "w") as f:
        f.write("This is a sample TXT document. It talks about technology and innovation. The future is bright.")
    
    # Create empty TXT
    with open(empty_txt_path, "w") as f:
        f.write("")

    # Create long TXT for truncation test
    long_text_content = "This is a very long text. " * 500 # 10000 characters
    with open(long_txt_path, "w") as f:
        f.write(long_text_content)

    test_user_free = sys.modules['utils.user_manager']._mock_users["mock_free_token"]['user_id']
    test_user_pro = sys.modules['utils.user_manager']._mock_users["mock_pro_token"]['user_id']
    test_user_premium = sys.modules['utils.user_manager']._mock_users["mock_premium_token"]['user_id']
    test_user_admin = sys.modules['utils.user_manager']._mock_users["mock_admin_token"]['user_id']

    print("\n--- Testing summarize_document function ---")

    # Test 1: Pro user, TXT document
    print("\n--- Test 1: Pro user, TXT document ---")
    sys.modules['utils.user_manager']._current_mock_user = test_user_pro
    summary_txt = summarize_document(txt_path, user_token=test_user_pro)
    print(f"Summary of '{txt_path.name}' (Pro user): {summary_txt[:100]}...")
    assert "Mock summary" in summary_txt
    assert "Original content length: 90 characters." in summary_txt # Check mock LLM output
    print("Test 1 Passed.")

    # Test 2: Premium user, DOCX document
    print("\n--- Test 2: Premium user, DOCX document ---")
    sys.modules['utils.user_manager']._current_mock_user = test_user_premium
    summary_docx = summarize_document(docx_path, user_token=test_user_premium)
    print(f"Summary of '{docx_path.name}' (Premium user): {summary_docx[:100]}...")
    assert "Mock summary" in summary_docx
    assert "Original content length: 70 characters." in summary_docx
    print("Test 2 Passed.")

    # Test 3: Free user, summarization disabled
    print("\n--- Test 3: Free user, summarization disabled ---")
    sys.modules['utils.user_manager']._current_mock_user = test_user_free
    summary_free = summarize_document(txt_path, user_token=test_user_free)
    print(f"Summary of '{txt_path.name}' (Free user): {summary_free}")
    assert "Error: Document summarization is not enabled for your current tier." in summary_free
    print("Test 3 Passed.")

    # Test 4: Admin user, long TXT document (should be truncated to admin's max_input_chars)
    print("\n--- Test 4: Admin user, long TXT document (truncation test) ---")
    sys.modules['utils.user_manager']._current_mock_user = test_user_admin
    summary_long_txt = summarize_document(long_txt_path, user_token=test_user_admin)
    print(f"Summary of '{long_txt_path.name}' (Admin user): {summary_long_txt[:100]}...")
    assert "Mock summary" in summary_long_txt
    # Admin's max_input_chars is float('inf') in mock, but actual truncation happens before LLM.
    # The mock LLM will report the length of the *truncated* input it received.
    # The warning log for truncation should indicate the truncation to 50000 (mocked admin limit)
    assert "Original content length: 50000 characters." in summary_long_txt # Check mock LLM output
    print("Test 4 Passed.")

    # Test 5: Non-existent file
    print("\n--- Test 5: Non-existent file ---")
    sys.modules['utils.user_manager']._current_mock_user = test_user_pro
    non_existent_path = test_dir / "non_existent.pdf"
    summary_non_existent = summarize_document(non_existent_path, user_token=test_user_pro)
    print(f"Summary of '{non_existent_path.name}': {summary_non_existent}")
    assert "Error: Document not found" in summary_non_existent
    print("Test 5 Passed.")

    # Test 6: Empty document
    print("\n--- Test 6: Empty document ---")
    sys.modules['utils.user_manager']._current_mock_user = test_user_pro
    summary_empty = summarize_document(empty_txt_path, user_token=test_user_pro)
    print(f"Summary of '{empty_txt_path.name}': {summary_empty}")
    assert "Error: Could not extract any readable text from the document." in summary_empty
    print("Test 6 Passed.")

    # Test 7: Unsupported file type
    print("\n--- Test 7: Unsupported file type ---")
    sys.modules['utils.user_manager']._current_mock_user = test_user_pro
    unsupported_path = test_dir / "image.jpg"
    with open(unsupported_path, "w") as f: # Create a dummy file
        f.write("dummy image content")
    summary_unsupported = summarize_document(unsupported_path, user_token=test_user_pro)
    print(f"Summary of '{unsupported_path.name}': {summary_unsupported}")
    assert "Error: Unsupported file type for summarization: .jpg" in summary_unsupported
    os.remove(unsupported_path) # Clean up dummy file
    print("Test 7 Passed.")

    print("\nAll summarize_document tests passed (mocked LLM and RBAC).")

    # Clean up dummy test directory
    if test_dir.exists():
        shutil.rmtree(test_dir)
        print(f"\nCleaned up test directory: {test_dir}")
